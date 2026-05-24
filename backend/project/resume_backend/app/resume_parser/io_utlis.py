import fitz  # PyMuPDF
import docx2txt
from docx import Document
from io import BytesIO



def extract_text_from_pdf(file_bytes):
    text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text


#__________________________________________________________________________
def extract_text_from_docx(file_bytes):
    # Save bytes to a temporary file for docx2txt
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    # Try docx2txt first
    text = docx2txt.process(tmp_path)
    # Clean up the temp file
    import os
    os.unlink(tmp_path)

    # If docx2txt found enough text, use it
    if text and len(text.strip()) > 30:  # adjust threshold as needed
        return text

    # Fallback: use python-docx extraction
    text_parts = []
    try:
        doc = Document(BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" ".join(row_text))
        for section in doc.sections:
            header = section.header
            footer = section.footer
            for para in header.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            for para in footer.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
    except Exception as e:
        print(f"[DOCX Extraction Error]: {e}")
    return "\n".join(text_parts)
