import fitz  # PyMuPDF
import docx
import io
from docx import Document
from io import BytesIO
from rapidfuzz import fuzz
import re
import spacy
from pyresparser import ResumeParser
import tempfile
import os
import nltk
import docx2txt
from difflib import get_close_matches



#_____________________________________________________________________________
def extract_text_from_pdf(file_bytes):
    text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text


#__________________________________________________________________________
def extract_text_from_docx(file_bytes):
    # Save bytes to a temporary file for docx2txt
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    # Try docx2txt first
    text = docx2txt.process(tmp_path)
    # Clean up the temp file
    import os
    os.unlink(tmp_path)

    # If docx2txt found enough text, use it
    if text and len(text.strip()) > 30:  # adjust threshold as needed
        return text

    # Fallback: use python-docx extraction
    text_parts = []
    try:
        doc = Document(BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" ".join(row_text))
        for section in doc.sections:
            header = section.header
            footer = section.footer
            for para in header.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            for para in footer.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
    except Exception as e:
        print(f"[DOCX Extraction Error]: {e}")
    return "\n".join(text_parts)


#_____________________________________________________________________________




# ------------------------ Personal Info Extraction -------------------------------------------------

#________________________________________email extraction________________________
def extract_email(text):
    email_match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    return email_match.group() if email_match else None



#________________________________________phone extraction_______________________
def extract_phone(text):
    phone_match = re.search(
        r'(\+?\d{1,3}[-.\s]?)?(\(?\d{2,4}\)?[-.\s]?){2,4}\d{3,4}', text)
    return phone_match.group() if phone_match else None



#________________________________________name extraction_________________________
def extract_name(text):
    lines = [line.strip() for line in text.strip().split('\n') if line.strip()]
    # Only look at the first 10 lines
    lines = lines[:10]
    skip_keywords = {"resume", "curriculum", "skills", "objective", "summary", "profile", "contact", "email", "phone", "address"}
    for line in lines:
        if line.endswith(":"):
            continue
        if any(kw in line.lower() for kw in skip_keywords):
            continue
        words = line.strip().split()
        if (
            len(words) >= 2
            and all(word.istitle() or word.isupper() for word in words)
            and not any(char.isdigit() for char in line)
            and len(line) < 40
        ):
            return " ".join(words)
    return None

#------------------------------------------------------
def fallback_extract_name(text):
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return None
#-----------------------------------------------------------



#_________________________address extraction_____________________________________
def extract_address(text):
    # Prefer lines that look like an address
    for line in text.split('\n'):
        if line.strip().lower().startswith("address:"):
            return line.strip()
        if re.search(r"\d{1,5} .+(road|street|avenue|block|sector|lane|colony|nagar|park)", line, re.IGNORECASE):
            return line.strip()
    # Fallback to spaCy NER
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(text)
    locations = [ent.text for ent in doc.ents if ent.label_ in ["GPE", "LOC"]]
    for loc in locations:
        if any(char.isdigit() for char in loc) or ',' in loc:
            return loc
    return locations[0] if locations else None



# _______________________________________Analyzing sections separately_________________
#------------------------------------------------------------------------------
def normalize_header(text):
    return re.sub(r'\s+', '', text).lower()

# All known headers normalized
SECTION_HEADERS = [
    "education", "experience", "work history", "skills", "professional skills", "professional skills and experience",
    "technical skills", "technical knowledge and skills", "key skills", "core competencies", "knowledge", "areas of expertise",
    "computer skills", "applications", "languages", "programming languages", "software", "tools", "technologies",
    "it skills", "it competencies", "it knowledge", "it expertise", "it capabilities", "it proficiencies", "it abilities",
    "it competencies and experience", "projects", "certifications", "summary", "objective", "activities", "interests",
    "contact", "references", "community service", "achievements", "awards", "publications", "volunteer work", "hobbies",
    "training", "courses", "professional development", "academic", "qualifications", "degree", "education history",
    "employee history", "employment history", "career history", "career summary", "career objective", "profile",
    "personal statement", "self-summary", "about me", "bio", "introduction", "soft skills"
]
NORMALIZED_HEADERS = set(normalize_header(h) for h in SECTION_HEADERS)

def extract_sections(resume_text):
    lines = resume_text.splitlines()
    sections = {}
    current_section = None
    buffer = []

    for line in lines:
        line_clean = line.strip().rstrip(":").strip()
        if not line_clean:
            continue

        normalized = normalize_header(line_clean)
        if normalized in NORMALIZED_HEADERS:
            # Save the current section before moving to next
            if current_section and buffer:
                sections[current_section] = '\n'.join(buffer).strip()
                buffer = []
            current_section = normalized
        elif current_section:
            buffer.append(line_clean)

    # Save last section
    if current_section and buffer:
        sections[current_section] = '\n'.join(buffer).strip()

    return sections



#_____________________________extracting education and experience___________________

# def extract_education(text):
#     lines = [line.strip() for line in text.split('\n')]
#     section_names = ["education", "academic", "qualification", "degree"]
#     block = find_section_block(lines, section_names)
#     if not block:
#         return None
#     return "\n".join(block)


# def extract_experience(text):
#     lines = [line.strip() for line in text.split('\n')]
#     section_names = ["experience", "employment", "work history", "professional experience"]
#     block = find_section_block(lines, section_names)
#     if not block:
#         return None

#     experience = []
#     for i, line in enumerate(block):
#         if (
#             re.search(r'\b(19|20)\d{2}\b', line) or
#             re.search(r'present', line, re.IGNORECASE) or
#             '|' in line or
#             (line == line.title() and 2 <= len(line.split()) <= 8) or
#             (line.isupper() and 2 <= len(line.split()) <= 8)
#         ):
#             experience.append(line)
#         elif i+1 < len(block) and re.search(r'\b(19|20)\d{2}\b', block[i+1]):
#             experience.append(line)
#     experience = list(dict.fromkeys(experience))
#     return "\n".join(experience) if experience else None




#_______________________________extracting skills from pdf___________________________________________________________________________________________

def extract_all_resume_skills(sections):
    import re

    def normalize_header(text):
        return re.sub(r'\s+', '', text.strip().lower())

    skill_section_keys = [
        "skills", "professional skills", "technical skills", "key skills", "core competencies", "knowledge", "areas of expertise",
        "computer skills", "professional skills and experience", "technical knowledge and skills", "applications", "languages",
        "programming languages", "software", "tools", "technologies", "it skills", "it competencies", "it knowledge",
        "it expertise", "it capabilities", "it proficiencies", "it abilities", "it competencies and experience", "soft skills"
    ]

    degree_keywords = {
        "bachelor", "master", "phd", "mba", "degree", "university", "college", "school", "gpa", "graduated"
    }
    non_skill_keywords = {
        "page", "available", "intern", "assistant", "manager", "technician", "coordinator", "volunteer", "contributor",
        "dean", "scholar", "project", "lab", "office", "university", "school", "college", "committee",
        "lead", "supervisor", "director", "secretary", "address", "dates", "contact", "phone", "email"
    }

    normalized_sections = {normalize_header(k): v for k, v in sections.items()}
    skill_results = {}
    global_seen = set()  # Track all skills seen so far

    for k in skill_section_keys:
        norm_k = normalize_header(k)
        if norm_k in normalized_sections:
            section_text = normalized_sections[norm_k]
            # Split by bullet, middot, comma, or semicolon
            raw_skills = re.split(r"[•·,;]", section_text)
            filtered_skills = []
            for s in raw_skills:
                s_clean = s.strip()
                s_lower = s_clean.lower()
                if (
                    s_clean
                    and len(s_clean) > 2
                    and s_lower not in degree_keywords
                    and s_lower not in non_skill_keywords
                    and not s_lower.isdigit()
                    and s_lower not in global_seen
                ):
                    filtered_skills.append(s_clean)
                    global_seen.add(s_lower)
            pretty_heading = k.title()
            skill_results[pretty_heading] = filtered_skills

    return skill_results

#_______________________extracting education and experience from pdf_______________________-
def clean_education(text):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    filtered_lines = []
    edu_keywords = [
        "bachelor", "master", "phd", "b.sc", "m.sc", "mba", "ba", "bs", "b.tech", "m.tech",
        "bca", "mca", "university", "college", "school", "institute", "arts", "science", "fine arts", "degree"
    ]
    date_pattern = re.compile(
        r"\b(?:(?:JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)[A-Z]*\s+)?(?:20XX|19\d{2}|20\d{2})"
        r"(?:\s*[–\-]\s*(?:PRESENT|20XX|19\d{2}|20\d{2}|(?:JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)[A-Z]*\s+\d{4}))?",
        re.IGNORECASE
    )
    for i, line in enumerate(lines):
        l = line.lower()
        # Only keep lines that look like degree, school, or year (sideheadings)
        if (
            any(kw in l for kw in edu_keywords)
            or date_pattern.search(line)
            or (1 <= len(line.split()) <= 7 and (line == line.title() or line.isupper()))
        ):
            filtered_lines.append(line)
        # If the next line is a date, treat this as a heading too
        elif i + 1 < len(lines) and date_pattern.search(lines[i + 1]):
            # Only include if this line is short (likely a heading)
            if 1 <= len(line.split()) <= 7:
                filtered_lines.append(line)
    # Remove duplicates, preserve order
    filtered_lines = list(dict.fromkeys(filtered_lines))
    return "\n".join(filtered_lines) if filtered_lines else None


def get_merged_section(sections, keys):
    """Merge multiple possible section contents into one string."""
    merged = []
    for k in keys:
        content = sections.get(k, "")
        if content:
            merged.append(content)
    return "\n\n".join(merged).strip() if merged else None



# ----------------------------------------------------------------------------
def clean_experience(text):
    import re
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if len(lines) > 12:
        # Fussy logic (your current code)
        filtered_lines = []
        section_headers = [
            "education", "experience", "work history", "professional experience", "skills", "projects", "certifications",
            "summary", "objective", "activities", "interests", "languages", "contact", "references", "community service",
            "achievements", "awards", "publications", "volunteer work", "hobbies", "training", "courses", "professional developoment",
            "academic", "qualifications", "degree", "education history", "employee history", "employment history", "career history",
            "career summary", "career objective", "profile", "personal statement", "self-summary", "about me", "bio", "intoduction",
            "work", "professional", "career", "employment", "experience summary"
        ]
        normalized_headers = set(re.sub(r'\s+', '', h).lower() for h in section_headers)
        date_pattern = re.compile(
            r"\b(?:(?:JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)[A-Z]*\s+)?(?:20XX|19\d{2}|20\d{2})"
            r"(?:\s*[–\-to]+\s*(?:PRESENT|20XX|19\d{2}|20\d{2}|present|(?:JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)[A-Z]*\s+\d{4}))?",
            re.IGNORECASE
        )
        for i, line in enumerate(lines):
            norm = re.sub(r'\s+', '', line).lower()
            if norm in normalized_headers and i != 0:
                break
            if (
                date_pattern.search(line)
                or (2 <= len(line.split()) <= 10 and (line.istitle() or line.isupper()) and not re.search(r"\d{4}", line))
            ):
                filtered_lines.append(line)
            elif i + 1 < len(lines) and date_pattern.search(lines[i + 1]):
                if 2 <= len(line.split()) <= 10 and (line.istitle() or line.isupper()):
                    filtered_lines.append(line)
        filtered_lines = list(dict.fromkeys(filtered_lines))
        return "\n".join(filtered_lines) if filtered_lines else None
    else:
        return normal_experience(lines)

def normal_experience(lines):
    import re
    section_headers = [
        "education", "experience", "work history", "professional experience", "skills", "projects", "certifications",
        "summary", "objective", "activities", "interests", "languages", "contact", "references", "community service",
        "achievements", "awards", "publications", "volunteer work", "hobbies", "training", "courses", "professional developoment",
        "academic", "qualifications", "degree", "education history", "employee history", "employment history", "career history",
        "career summary", "career objective", "profile", "personal statement", "self-summary", "about me", "bio", "intoduction",
        "work", "professional", "career", "employment", "experience summary"
    ]
    normalized_headers = set(re.sub(r'\s+', '', h).lower() for h in section_headers)
    date_pattern = re.compile(
        r"\b(?:(?:JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)[A-Z]*\s+)?(?:20XX|19\d{2}|20\d{2})"
        r"(?:\s*[–\-to]+\s*(?:PRESENT|20XX|19\d{2}|20\d{2}|present|(?:JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)[A-Z]*\s+\d{4}))?",
        re.IGNORECASE
    )
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]
        norm = re.sub(r'\s+', '', line).lower()
        # Only start a block if this is a heading (not a section header)
        if (2 <= len(line.split()) <= 15 and norm not in normalized_headers):
            # If date is on the same line, split and add both
            match = date_pattern.search(line)
            if match:
                date_str = match.group()
                heading = line.replace(date_str, '').strip(" ,.-")
                if heading:
                    result.append(heading)
                result.append(date_str)
            # If next line is a date, add it
            elif i + 1 < len(lines) and date_pattern.search(lines[i + 1]):
                result.append(line)
                result.append(lines[i + 1])
                i += 1
            else:
                result.append(line)
        i += 1
    # Remove duplicates, preserve order
    result = list(dict.fromkeys(result))
    return "\n".join(result) if result else None


#_________________________extractiing projects_______________________________________
def extract_projects(sections):
    import re

    project_section_keys = [
        "projects", "personalprojects", "academicprojects", "keyprojects", "majorprojects", "projectexperience",
        "notableprojects", "selectedprojects", "relevantprojects", "projectwork", 
    ]
    ignore_headings = {
        "sample resume", "curriculum vitae", "resume", "cv", "profile", "page", "contact", "summary"
    }
    normalized_sections = {re.sub(r'\s+', '', k).lower(): v for k, v in sections.items()}

    # Find the first matching project section
    projects_text = ""
    for k in project_section_keys:
        if k in normalized_sections:
            projects_text = normalized_sections[k]
            break
    if not projects_text.strip():
        return []

    lines = [line.strip() for line in projects_text.splitlines() if line.strip()]
    projects = []
    i = 0
    while i < len(lines):
        line = lines[i]
        # Ignore lines in ignore_headings
        if line.strip().lower() in ignore_headings:
            i += 1
            continue

        # Heading: any non-bullet, non-empty line not in ignore list
        is_heading = (
            not re.match(r"^[-*•·]\s*", line) and
            not line.strip().lower() in ignore_headings
        )

        if is_heading:
            project = {"name": line}
            desc_lines = []
            j = i + 1
            while j < len(lines):
                desc_candidate = lines[j]
                # Stop at next heading (non-bullet, not in ignore list)
                next_is_heading = (
                    not re.match(r"^[-*•·]\s*", desc_candidate) and
                    not desc_candidate.strip().lower() in ignore_headings
                )
                if next_is_heading:
                    break
                # Accept bullet/dash or plain lines as description
                if re.match(r"^[-*•·]\s*", desc_candidate):
                    desc_line = re.sub(r"^[-*•·]\s*", "", desc_candidate)
                    desc_lines.append(desc_line)
                elif desc_candidate:
                    desc_lines.append(desc_candidate)
                j += 1

            if desc_lines:
                description = " ".join(desc_lines)
                project["description"] = description

                # Extract date/duration
                date_pattern = re.compile(
                    r"((Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}\s*[–\-to]+\s*(Present|\d{4})|"
                    r"\d{4}\s*[–\-to]+\s*(Present|\d{4})|"
                    r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})",
                    re.IGNORECASE
                )
                date_match = date_pattern.search(description)
                project["date"] = date_match.group().strip() if date_match else None

                # Extract first link (http/https)
                link_match = re.search(r"https?://\S+", description)
                project["link"] = link_match.group().strip() if link_match else None
            else:
                project["description"] = None
                project["date"] = None
                project["link"] = None

            projects.append(project)
            i = j
        else:
            i += 1
    return projects




#__________________extracting education, experience and skills from docx________________________

def docx_clean_education(lines):
    result = []
    block = []
    for i, line in enumerate(lines):
        l = line.lower()
        if re.search(r"(university|institute|college)", l):
            if block:
                result.append("\n".join(block))
                block = []
            block.append(line)
        elif re.search(r"(bachelor|master|phd|mba|arts|science|fine arts|degree)", l):
            block.append(line)
        elif re.search(r"(20XX|19\d{2}|20\d{2}|present)", line, re.IGNORECASE):
            block.insert(1, line)
    if block:
        result.append("\n".join(block))
    return "\n\n".join(result).strip() if result else None




#--------------------------------------------------------------------------------------------
def docx_clean_skills(lines):
    import re
    skills = []
    # Keywords to exclude (not skills)
    exclude_keywords = [
        "university", "college", "school", "institute", "academy", "bachelor", "master", "phd", "mba", "degree",
        "company", "corporation", "inc", "ltd", "llc", "solutions", "technologies", "systems", "pvt", "limited",
        "experience", "project", "resume", "curriculum", "vitae", "profile", "summary", "objective", "history"
    ]
    for line in lines:
        # Only accept bullet points or short lines that look like skills
        if re.match(r"^[-*•·]\s*", line):
            skill = re.sub(r"^[-*•·]\s*", "", line).strip()
        else:
            skill = line.strip()
        # Exclude lines with years/dates, numbers, excluded keywords, or too long/short
        if (
            skill
            and 1 <= len(skill.split()) <= 5
            and not re.search(r"\d{4}", skill)
            and not any(kw in skill.lower() for kw in exclude_keywords)
            and not re.search(r"\d", skill)
            and not skill.isupper()
            and not skill.istitle()
        ):
            skills.append(skill)
    # Remove duplicates, preserve order
    skills = list(dict.fromkeys(skills))
    return skills

def extract_skills_from_docx(sections):
    # Try modern/table/cell extraction first
    docx_skills = docx_clean_skills([line.strip() for line in sections.get("skills", "").splitlines() if line.strip()])
    # If not enough skills, or empty, fallback to PDF-style extraction
    if not docx_skills or len(docx_skills) < 3:
        # Use the robust PDF-style skill extraction
        skill_dict = extract_all_resume_skills(sections)
        # Flatten all skill lists into one, removing duplicates
        all_skills = []
        seen = set()
        for skill_list in skill_dict.values():
            for skill in skill_list:
                if skill not in seen:
                    all_skills.append(skill)
                    seen.add(skill)
        return all_skills
    return docx_skills


#--------------------------------------------------------------------------------------------------
def normalize_header(text):
    return re.sub(r'\s+', '', text.strip().lower())

def extract_experience_from_docx_tables(file_bytes):

    section_headers = [
        "education", "skills", "projects", "certifications", "summary", "objective", "activities", "interests",
        "languages", "contact", "references", "community service", "achievements", "awards", "publications",
        "volunteer work", "hobbies", "training", "courses", "profile"
    ]
    normalized_headers = set(normalize_header(h) for h in section_headers)
    date_pattern = re.compile(r"(20\d{2}|19\d{2})(\s*[–\-to]+\s*(Present|20\d{2}|19\d{2}|present))?", re.IGNORECASE)

    try:
        doc = Document(BytesIO(file_bytes))
        for table in doc.tables:
            cells = []
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        cells.append(text)
            # Find the index of the "EXPERIENCE" cell
            exp_idx = None
            for idx, cell in enumerate(cells):
                if normalize_header(cell) == "experience":
                    exp_idx = idx
                    break
            if exp_idx is not None:
                i = exp_idx + 1
                blocks = []
                while i < len(cells):
                    if normalize_header(cells[i]) in normalized_headers:
                        break
                    # Company
                    company = cells[i]
                    i += 1
                    # Date (optional)
                    date = ""
                    if i < len(cells) and date_pattern.match(cells[i]):
                        date = cells[i]
                        i += 1
                    # Description (up to 5 lines)
                    desc_lines = []
                    desc_count = 0
                    while i < len(cells) and desc_count < 5:
                        if normalize_header(cells[i]) in normalized_headers:
                            break
                        # Stop if next cell looks like a company or date
                        if (2 <= len(cells[i].split()) <= 8 and (cells[i].istitle() or cells[i].isupper()) and not re.search(r"\d{4}", cells[i])):
                            break
                        if date_pattern.match(cells[i]):
                            break
                        desc_lines.append(cells[i])
                        desc_count += 1
                        i += 1
                    block = [company]
                    if date:
                        block.append(date)
                    if desc_lines:
                        block.extend(desc_lines)
                    blocks.append("\n".join(block))
                if blocks:
                    return "\n\n".join(blocks)
    except Exception as e:
        print(f"[DOCX Table Experience Extraction Error]: {e}")
    return None



def extract_experience_from_docx_paragraphs(sections):
    import re
    date_pattern = re.compile(r"(20\d{2}|19\d{2})(\s*[–\-to]+\s*(Present|20\d{2}|19\d{2}|present))?", re.IGNORECASE)
    section_headers = [
        "education", "skills", "projects", "certifications", "summary", "objective", "activities", "interests",
        "languages", "contact", "references", "community service", "achievements", "awards", "publications",
        "volunteer work", "hobbies", "training", "courses", "profile"
    ]
    normalized_headers = set(normalize_header(h) for h in section_headers)
    exp_keys = [
        "experience", "workhistory", "professionalexperience", "employmenthistory",
        "careerhistory", "careerexperience", "workexperience", "jobhistory", "jobexperience"
    ]
    exp_section = None
    for k in exp_keys:
        if k in sections:
            exp_section = sections[k]
            break
    if not exp_section:
        for k, v in sections.items():
            if "experience" in k:
                exp_section = v
                break
    if not exp_section:
        return ""
    lines = [line.strip() for line in exp_section.splitlines() if line.strip()]
    blocks = []
    i = 0
    while i < len(lines):
        # Company
        if (2 <= len(lines[i].split()) <= 8 and (lines[i].istitle() or lines[i].isupper()) and not re.search(r"\d{4}", lines[i]) and normalize_header(lines[i]) not in normalized_headers):
            company = lines[i]
            i += 1
            # Date (optional)
            date = ""
            if i < len(lines) and date_pattern.match(lines[i]):
                date = lines[i]
                i += 1
            # Description (up to 5 lines)
            desc_lines = []
            desc_count = 0
            while i < len(lines) and desc_count < 5:
                if normalize_header(lines[i]) in normalized_headers:
                    break
                if (2 <= len(lines[i].split()) <= 8 and (lines[i].istitle() or lines[i].isupper()) and not re.search(r"\d{4}", lines[i])):
                    break
                if date_pattern.match(lines[i]):
                    break
                desc_lines.append(lines[i])
                desc_count += 1
                i += 1
            block = [company]
            if date:
                block.append(date)
            if desc_lines:
                block.extend(desc_lines)
            blocks.append("\n".join(block))
        else:
            i += 1
    return "\n\n".join(blocks)


def extract_experience_block_from_sections(sections):
    # Try all normalized experience-related keys
    exp_keys = [
        "experience", "workhistory", "professionalexperience", "employmenthistory",
        "careerhistory", "careerexperience", "workexperience", "jobhistory", "jobexperience"
    ]
    for k in exp_keys:
        if k in sections:
            return sections[k]
    # Fallback: any section containing "experience"
    for k, v in sections.items():
        if "experience" in k:
            return v
    return ""

def clean_experience_block(text):
    import re
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    blocks = []
    i = 0
    date_pattern = re.compile(r"(20\d{2}|19\d{2})(\s*[–\-to]+\s*(Present|20\d{2}|19\d{2}|present))?", re.IGNORECASE)
    while i < len(lines):
        # Company
        if (2 <= len(lines[i].split()) <= 8 and (lines[i].istitle() or lines[i].isupper())):
            company = lines[i]
            i += 1
            # Date (optional)
            date = ""
            if i < len(lines) and date_pattern.match(lines[i]):
                date = lines[i]
                i += 1
            # Description (up to 5 lines)
            desc_lines = []
            desc_count = 0
            while i < len(lines) and desc_count < 5:
                # Stop if next line looks like a new company or date
                if (2 <= len(lines[i].split()) <= 8 and (lines[i].istitle() or lines[i].isupper())):
                    break
                if date_pattern.match(lines[i]):
                    break
                desc_lines.append(lines[i])
                desc_count += 1
                i += 1
            block = [company]
            if date:
                block.append(date)
            if desc_lines:
                block.extend(desc_lines)
            blocks.append("\n".join(block))
        else:
            i += 1
    return "\n\n".join(blocks)


def extract_experience_blocks_flexible(raw_text):
    import re
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    section_headers = [
        "education", "projects", "certifications", "summary", "objective", "activities", "interests",
        "languages", "contact", "references", "community service", "achievements", "awards", "publications",
        "volunteer work", "hobbies", "training", "courses", "profile"
    ]
    normalized_headers = set(h.replace(" ", "").lower() for h in section_headers)
    # Find the EXPERIENCE section
    start = None
    for idx, line in enumerate(lines):
        if line.replace(" ", "").lower() == "experience":
            start = idx + 1
            break
    if start is None:
        return ""
    # Collect lines until next major section header (not skills)
    exp_lines = []
    for line in lines[start:]:
        norm = line.replace(" ", "").lower()
        if norm in normalized_headers:
            break
        exp_lines.append(line)
    date_pattern = re.compile(r"(20\d{2}|19\d{2})(\s*[–\-to]+\s*(Present|20\d{2}|19\d{2}|present))?", re.IGNORECASE)
    skip_headers = {"skills"}
    skill_keywords = {"design", "python", "java", "c++", "django", "machine learning", "communication", "problem solving"}
    blocks = []
    i = 0
    company = None
    date = ""
    desc_lines = []
    while i < len(exp_lines):
        line = exp_lines[i]
        norm = line.replace(" ", "").lower()
        words = line.split()
        is_skill_line = (
            norm in skip_headers or
            (1 <= len(words) <= 5 and any(kw in line.lower() for kw in skill_keywords) and not date_pattern.match(line))
        )
        # Skip skill lines
        if is_skill_line:
            i += 1
            continue
        # Company
        is_company = (2 <= len(words) <= 8 and (line.istitle() or line.isupper()))
        if is_company:
            if company and (desc_lines or date):
                block = [company]
                if date:
                    block.append(date)
                if desc_lines:
                    block.extend(desc_lines)
                blocks.append("\n".join(block))
                date = ""
                desc_lines = []
            company = line
            i += 1
            # Date (optional)
            if i < len(exp_lines) and date_pattern.match(exp_lines[i]):
                date = exp_lines[i]
                i += 1
            else:
                date = ""
            # Description (up to 5 lines)
            desc_lines = []
            desc_count = 0
            while i < len(exp_lines) and desc_count < 5:
                desc_line = exp_lines[i]
                norm2 = desc_line.replace(" ", "").lower()
                words2 = desc_line.split()
                is_skill_line2 = (
                    norm2 in skip_headers or
                    (1 <= len(words2) <= 5 and any(kw in desc_line.lower() for kw in skill_keywords) and not date_pattern.match(desc_line))
                )
                if is_skill_line2:
                    i += 1
                    continue
                if (2 <= len(words2) <= 8 and (desc_line.istitle() or desc_line.isupper())):
                    break
                if date_pattern.match(desc_line):
                    break
                desc_lines.append(desc_line)
                desc_count += 1
                i += 1
        else:
            i += 1
    # Add the last block if any
    if company and (desc_lines or date):
        block = [company]
        if date:
            block.append(date)
        if desc_lines:
            block.extend(desc_lines)
        blocks.append("\n".join(block))
    return "\n\n".join(blocks)

#__________________________Main function___________________________________

def parse_resume(filename, file_bytes):
    if filename.endswith(".pdf"):
        raw_text = extract_text_from_pdf(file_bytes)
        sections = extract_sections(raw_text)
        education = clean_education(sections.get("education", ""))
        # Merge all experience-related sections

        
        
        experience_section_keys = [
        "experience", "work history", "professional experience", "employment history",
        "career history", "career experience", "work experience", "job history", "job experience",
        "work", "professional", "career", "employment", "experience summary",
        "career summary", "career objective"
        ]
        normalized_keys = [normalize_header(k) for k in experience_section_keys]
        experience_text = get_merged_section(sections, normalized_keys)
        experience = clean_experience(experience_text or "")
        skills = extract_all_resume_skills(sections)



    #_________________________________________________________________________________________
    elif filename.endswith(".docx"):
        raw_text = extract_text_from_docx(file_bytes)
        sections = extract_sections(raw_text)
        education = docx_clean_education([line.strip() for line in sections.get("education", "").splitlines() if line.strip()])
        # For skills, use the improved function
        skills = extract_skills_from_docx(sections)
       # 1. Try table extraction
        experience = extract_experience_from_docx_tables(file_bytes)
        # 2. Fallback: extract from paragraphs/sideheadings
        if not experience:
            exp_block = extract_experience_block_from_sections(sections)
            experience = clean_experience_block(exp_block)
        # 3. Fallback: extract from raw text (handles messy headers like SKILLS between jobs)
        if not experience:
            experience = extract_experience_blocks_flexible(raw_text)



    #_______________________________________________________________________________________
    name = extract_name(raw_text)
    if not name:
        name = fallback_extract_name(raw_text)
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    address = extract_address(raw_text)
    # education = extract_education(raw_text)
    # experience = extract_experience(raw_text)
    
    projects = extract_projects(sections)
    

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "address": address,
        "education": education,
        "experience": experience,
        # "education": education,
        # "experience": experience,
        "skills_found": skills,
        "projects": projects,
        "resume_text": raw_text# limit output for preview
    }
